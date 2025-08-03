"""
Document Loader Module
Handles loading and preprocessing of various document formats (PDF, DOCX, TXT)
"""

import os
import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF
import docx
from dataclasses import dataclass

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Document container with metadata"""
    content: str
    filename: str
    file_type: str
    page_count: Optional[int] = None
    metadata: Optional[Dict[str, Any]] = None

class DocumentLoader:
    """
    Loads documents from various file formats and extracts text content
    """
    
    def __init__(self, docs_path: str = "docs/"):
        self.docs_path = Path(docs_path)
        self.supported_formats = {'.pdf', '.docx', '.txt'}
        
    def load_document(self, file_path: str) -> Document:
        """
        Load a single document and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document object with extracted content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        logger.info(f"Loading document: {file_path.name}")
        
        try:
            if file_ext == '.pdf':
                return self._load_pdf(file_path)
            elif file_ext == '.docx':
                return self._load_docx(file_path)
            elif file_ext == '.txt':
                return self._load_txt(file_path)
        except Exception as e:
            logger.error(f"Error loading {file_path}: {str(e)}")
            raise
    
    def load_all_documents(self) -> List[Document]:
        """
        Load all supported documents from the docs directory
        
        Returns:
            List of Document objects
        """
        documents = []
        
        if not self.docs_path.exists():
            logger.warning(f"Documents directory not found: {self.docs_path}")
            return documents
        
        for file_path in self.docs_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    doc = self.load_document(file_path)
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {str(e)}")
                    continue
        
        logger.info(f"Loaded {len(documents)} documents")
        return documents
    
    def _load_pdf(self, file_path: Path) -> Document:
        """Extract text from PDF file, with preprocessing to remove headers/footers."""
        content = ""
        page_count = 0
        
        with fitz.open(file_path) as pdf_doc:
            page_count = len(pdf_doc)
            
            for page_num in range(page_count):
                page = pdf_doc[page_num]
                full_text = page.get_text("text")
                
                # Simple heuristic to remove headers and footers
                # Assumes header is in top 10% and footer is in bottom 10% of the page
                lines = full_text.split('\n')
                meaningful_lines = lines[int(len(lines)*0.1) : int(len(lines)*0.9)]
                content += "\n".join(meaningful_lines) + "\n"
        
        return Document(
            content=self._post_process_text(content),
            filename=file_path.name,
            file_type="pdf",
            page_count=page_count,
            metadata={"pages": page_count}
        )

    def _post_process_text(self, text: str) -> str:
        """Clean up extracted text."""
        # Remove multiple consecutive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove page numbers and other small, isolated text fragments
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        return text.strip()
    
    def _load_docx(self, file_path: Path) -> Document:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        content = ""
        
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + " "
                content += "\n"
        
        return Document(
            content=content.strip(),
            filename=file_path.name,
            file_type="docx",
            metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
        )
    
    def _load_txt(self, file_path: Path) -> Document:
        """Load text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return Document(
            content=content.strip(),
            filename=file_path.name,
            file_type="txt",
            metadata={"size": len(content)}
        )
    
    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """
        Get statistics about loaded documents
        
        Args:
            documents: List of Document objects
            
        Returns:
            Dictionary with document statistics
        """
        if not documents:
            return {"total_docs": 0}
        
        stats = {
            "total_docs": len(documents),
            "file_types": {},
            "total_content_length": 0,
            "avg_content_length": 0
        }
        
        for doc in documents:
            # Count file types
            if doc.file_type in stats["file_types"]:
                stats["file_types"][doc.file_type] += 1
            else:
                stats["file_types"][doc.file_type] = 1
            
            # Calculate content lengths
            stats["total_content_length"] += len(doc.content)
        
        stats["avg_content_length"] = stats["total_content_length"] / len(documents)
        
        return stats

# Example usage and testing
if __name__ == "__main__":
    loader = DocumentLoader("docs/")
    
    # Load all documents
    docs = loader.load_all_documents()
    
    # Print statistics
    stats = loader.get_document_stats(docs)
    print("Document Statistics:")
    print(f"Total documents: {stats['total_docs']}")
    print(f"File types: {stats['file_types']}")
    print(f"Average content length: {stats['avg_content_length']:.0f} characters")
    
    # Print first few characters of each document
    for doc in docs[:3]:  # Show first 3 documents
        print(f"\n--- {doc.filename} ({doc.file_type}) ---")
        print(doc.content[:200] + "..." if len(doc.content) > 200 else doc.content)