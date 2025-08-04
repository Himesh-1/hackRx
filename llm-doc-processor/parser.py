
"""
Parser Module
This module is responsible for parsing different file types.
"""

import fitz  # PyMuPDF
import docx
import logging
import re
from pathlib import Path
from typing import Dict, Any, Optional
from dataclasses import dataclass
import email
from email import policy
from email.parser import BytesParser

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Document container with metadata"""
    content: str
    filename: str
    file_type: str
    page_count: Optional[int] = None
    metadata: Optional[Dict[str, Any]] = None

class DocumentParser:
    """
    Parses various document formats and extracts text content.
    """

    def _post_process_text(self, text: str) -> str:
        """
        Cleans up extracted text by removing excessive newlines and isolated page numbers.

        Args:
            text (str): The raw text extracted from a document.

        Returns:
            str: The cleaned text.
        """
        # Remove multiple consecutive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove page numbers and other small, isolated text fragments
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        return text.strip()

    def parse_pdf(self, file_path: Path) -> Document:
        """
        Extracts text from a PDF file, applying preprocessing to remove headers/footers.

        Args:
            file_path (Path): The path to the PDF file.

        Returns:
            Document: A Document object containing the extracted text and metadata.
        """
        content = ""
        page_count = 0

        with fitz.open(file_path) as pdf_doc:
            page_count = len(pdf_doc)

            for page_num in range(page_count):
                page = pdf_doc[page_num]
                full_text = page.get_text("text")

                # Simple heuristic to remove headers and footers
                # Assumes header is in top 10% and footer is in bottom 10% of the page
                lines = full_text.split('\n')
                meaningful_lines = lines[int(len(lines)*0.1) : int(len(lines)*0.9)]
                content += "\n".join(meaningful_lines) + "\n"

        return Document(
            content=self._post_process_text(content),
            filename=file_path.name,
            file_type="pdf",
            page_count=page_count,
            metadata={"pages": page_count}
        )

    def parse_docx(self, file_path: Path) -> Document:
        """
        Extracts text from a DOCX file.

        Args:
            file_path (Path): The path to the DOCX file.

        Returns:
            Document: A Document object containing the extracted text and metadata.
        """
        doc = docx.Document(file_path)
        content = ""

        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + " "
                content += "\n"

        return Document(
            content=content.strip(),
            filename=file_path.name,
            file_type="docx",
            metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
        )

    def parse_txt(self, file_path: Path) -> Document:
        """
        Loads text from a TXT file.

        Args:
            file_path (Path): The path to the TXT file.

        Returns:
            Document: A Document object containing the extracted text and metadata.
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return Document(
            content=content.strip(),
            filename=file_path.name,
            file_type="txt",
            metadata={"size": len(content)}
        )

    def parse_email(self, file_path: Path) -> Document:
        """
        Extracts text from .eml or .msg email files.
        Requires 'extract-msg' for .msg files and Python's 'email' module for .eml.

        Args:
            file_path (Path): The path to the email file.

        Returns:
            Document: A Document object containing the extracted text and metadata.
        """
        try:
            import extract_msg
        except ImportError:
            logger.warning("extract-msg not found. Cannot process .msg files. Please install it: pip install extract-msg")
            extract_msg = None

        content = ""
        if file_path.suffix.lower() == '.msg':
            if extract_msg:
                try:
                    msg = extract_msg.Message(file_path)
                    content = f"From: {msg.sender}\nTo: {msg.to}\nSubject: {msg.subject}\n\n{msg.body}"
                except Exception as e:
                    logger.error(f"Error processing .msg file {file_path}: {e}")
                    content = ""
            else:
                raise ValueError("extract-msg library is required to process .msg files.")
        elif file_path.suffix.lower() == '.eml':
            with open(file_path, 'rb') as fp:
                msg = BytesParser(policy=policy.default).parse(fp)

            content += f"From: {msg['from']}\n"
            content += f"To: {msg['to']}\n"
            content += f"Subject: {msg['subject']}\n\n"

            if msg.is_multipart():
                for part in msg.walk():
                    ctype = part.get_content_type()
                    cdisp = str(part.get('Content-Disposition'))

                    # Look for plain text parts, but not attachments
                    if ctype == 'text/plain' and 'attachment' not in cdisp:
                        content += part.get_payload(decode=True).decode(errors='ignore')
                        break # Take the first plain text part
            else:
                content += msg.get_payload(decode=True).decode(errors='ignore')

        return Document(
            content=self._post_process_text(content),
            filename=file_path.name,
            file_type=file_path.suffix.lower().replace('.', ''),
            metadata={"source_file": str(file_path)}
        )
